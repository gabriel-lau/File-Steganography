"""
hello yes, erm you need to change the document paths under main if you wanna run this script separately.
this is before integration ^^ [for .docx files with LSB choice though idk why 0, 1 and 2 dont work 0-0]
"""
import os
from docx import Document


def encode_to_doc(base_file, payload, num_lsb):
    base_doc = Document(base_file)

    payload_bin = ''.join(format(ord(c), '08b') for c in payload)

    zero_char = '\u200B'  # zero-width space
    one_char = '\u200C'  # zero-width non-joiner
    boundary_char = '\u200D'  # zero-width joiner

    stego_text = boundary_char

    for i in range(len(payload_bin)):
        bit = payload_bin[i]
        if i % 3 < num_lsb:
            if bit == '0':
                stego_text += zero_char
            else:
                stego_text += one_char
        else:
            stego_text += bit

    base_doc_text = "\n".join([paragraph.text for paragraph in base_doc.paragraphs])
    stego_text = base_doc_text + stego_text

    stego_doc = Document()
    stego_doc.add_paragraph(stego_text)

    root, ext = os.path.splitext(base_file)
    dir_path = os.path.dirname(root)
    base_name = os.path.splitext(os.path.basename(base_file))[0]
    output = base_name + f"_encoded" + ext
    print('out_file_name:', output)
    print('directory_path', dir_path)
    encoded_file = os.path.join(dir_path, output)
    encoded_file = 'encoded_doc.docx'
    stego_doc.save(encoded_file)
    print("A very successful file creation has occurred :')")
    return encoded_file


def decode_frm_doc(stego_file, num_lsb):
    stego_doc = Document(stego_file)

    stego_text = "\n".join([paragraph.text for paragraph in stego_doc.paragraphs])

    find_start = stego_text.find('\u200D') + 1

    pay_load = stego_text[find_start:]

    payload_bin = ''
    for i in range(len(pay_load)):
        char = pay_load[i]
        if i % 3 < num_lsb:
            if char == '\u200B':
                payload_bin += '0'
            elif char == '\u200C':
                payload_bin += '1'

    decoded_text = ''
    for i in range(0, len(payload_bin), 8):
        byte = payload_bin[i:i + 8]
        if len(byte) == 8:
            try:
                decoded_text += chr(int(byte, 2))
            except ValueError:
                break

    return decoded_text


if __name__ == '__main__':
    payload = "jed was definitely here wakakakakaka"
    base_file = r"C:\Users\USER\PycharmProjects\stegosaurus\Docx\base_doc.docx"
    stego_file = r"C:\Users\USER\PycharmProjects\stegosaurus\Docx\base_doc_encoded.docx"

    num_lsb_to_use = 4  # (e.g., 0, 1, 2, ..., 5)
    encode_to_doc(base_file, payload, num_lsb_to_use)
    decoded_text = decode_frm_doc(stego_file, num_lsb_to_use)
    print("decoded thingy:", decoded_text)