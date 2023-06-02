"""
hello yes, erm you need to change the document paths under main if you wanna run this script separately.
this is before integration ^^ [for .docx files]
"""
import os
from docx import Document


def encode_to_doc(base_file, payload):
    base_doc = Document(base_file)

    payload_bin = ''.join(format(ord(c), '08b') for c in payload)

    zero_char = '\u200B'  # zero-width space
    one_char = '\u200C'  # zero-width non-joiner
    boundary_char = '\u200D'  # zero-width joiner

    stego_text = boundary_char

    for bit in payload_bin:
        if bit == '0':
            stego_text += zero_char
        else:
            stego_text += one_char

    base_doc_text = "\n".join([paragraph.text for paragraph in base_doc.paragraphs])
    stego_text = base_doc_text + stego_text

    stego_doc = Document()
    stego_doc.add_paragraph(stego_text)

    root, ext = os.path.splitext(base_file)
    dir_path = os.path.dirname(root)
    base_name = os.path.splitext(os.path.basename(base_file))[0]
    output = base_name + "_encoded" + ext
    print('out_file_name:', output)
    print('directory_path', dir_path)
    encoded_file = os.path.join(dir_path, output)
    stego_doc.save(encoded_file)
    print("a very successful file creation has occurred :')")


def decode_frm_doc(stego_file):
    stego_doc = Document(stego_file)

    stego_text = "\n".join([paragraph.text for paragraph in stego_doc.paragraphs])

    find_start = stego_text.find('\u200D') + 1

    pay_load = stego_text[find_start:]

    payload_bin = ''
    for char in pay_load:
        if char == '\u200B':
            payload_bin += '0'
        elif char == '\u200C':
            payload_bin += '1'

    decoded_text = ''
    for i in range(0, len(payload_bin), 8):
        byte = payload_bin[i:i + 8]
        decoded_text += chr(int(byte, 2))

    return decoded_text


if __name__ == '__main__':
    payload = "jed was definitely here wakakakakaka"
    base_file = r"C:\Users\USER\PycharmProjects\stegosaurus\Docx\base_doc.docx"
    stego_file = r"C:\Users\USER\PycharmProjects\stegosaurus\Docx\base_doc_encoded.docx"
    encode_to_doc(base_file, payload)
    decoded_text = decode_frm_doc(stego_file)
    print("decoded thingy:", decoded_text)
